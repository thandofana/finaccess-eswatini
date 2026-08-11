"""Build the professional Phase 1-13 implementation and rationale report."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "reports" / "project_documentation"
OUTPUT_FILE = OUTPUT_DIR / "FinAccess_Eswatini_Phases_1_to_13_Report.docx"

BLUE = "386BD1"
NAVY = "294A78"
TEAL = "2E9E92"
PALE_BLUE = "EAF1FB"
PALE_TEAL = "E5F4F2"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "697B91"
DARK = "263F61"
WHITE = "FFFFFF"
GOLD = "C49137"
RED = "A23B3B"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="D7E0EB", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=9, color=MID_GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(DARK)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.10

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("FINACCESS ESWATINI   |   PHASE 1-13 IMPLEMENTATION REPORT")
    set_run_font(r, size=8.5, color=MID_GRAY, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D9E3F0")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [6500, 2860], indent_dxa=0)
    table._tbl.tblPr.remove(table._tbl.tblPr.find(qn("w:tblBorders"))) if table._tbl.tblPr.find(qn("w:tblBorders")) is not None else None
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = left.add_run("Developed by Thando F. Dlamini")
    set_run_font(r, size=8.5, color=MID_GRAY)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right.add_run("Page ")
    set_run_font(r, size=8.5, color=MID_GRAY)
    add_page_field(right, "PAGE")
    r = right.add_run(" of ")
    set_run_font(r, size=8.5, color=MID_GRAY)
    add_page_field(right, "NUMPAGES")

    first_footer = section.first_page_footer
    fp = first_footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("Developed by Thando F. Dlamini")
    set_run_font(r, size=9, color=MID_GRAY)


def add_paragraph(doc, text="", *, bold_prefix=None, italic=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=color or DARK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, italic=italic, color=color or DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color or DARK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        if isinstance(item, tuple):
            label, detail = item
            r = p.add_run(label)
            set_run_font(r, bold=True, color=NAVY)
            r = p.add_run(detail)
            set_run_font(r, color=DARK)
        else:
            r = p.add_run(item)
            set_run_font(r, color=DARK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.keep_together = True
        r = p.add_run(item)
        set_run_font(r, color=DARK)


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label.upper())
    set_run_font(r, size=8.5, color=accent, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=DARK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_status_table(doc, status, purpose):
    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [1750, 7610])
    set_table_borders(table)
    for row_index, (label, value) in enumerate((('STATUS', status), ('PHASE PURPOSE', purpose))):
        label_cell, value_cell = table.rows[row_index].cells
        set_cell_shading(label_cell, PALE_BLUE)
        p = label_cell.paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, size=8.5, color=BLUE, bold=True)
        p = value_cell.paragraphs[0]
        r = p.add_run(value)
        set_run_font(r, size=10, color=DARK, bold=(row_index == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def set_picture_alt_text(inline_shape, description):
    """Attach accessibility text to an inline image in the generated DOCX."""
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_phase(doc, number, title, purpose, tasks, reasoning, validation, note=None, figure=None, figure_caption=None):
    doc.add_page_break()
    h = doc.add_heading(f"Phase {number}  |  {title}", level=1)
    h.paragraph_format.space_before = Pt(0)
    add_status_table(doc, "PASSED WITH NOTES", purpose)
    doc.add_heading("Tasks completed", level=2)
    add_bullets(doc, tasks)
    doc.add_heading("Reasoning behind the work", level=2)
    for label, text in reasoning:
        add_paragraph(doc, f"{label}: {text}", bold_prefix=f"{label}: ")
    doc.add_heading("Validation and handoff", level=2)
    add_bullets(doc, validation)
    if note:
        add_callout(doc, note[0], note[1], fill=note[2] if len(note) > 2 else PALE_BLUE, accent=note[3] if len(note) > 3 else BLUE)
    if figure and Path(figure).is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        picture = p.add_run().add_picture(str(figure), width=Inches(6.15))
        set_picture_alt_text(picture, figure_caption or "Validated project figure")
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(3)
        cap.paragraph_format.space_after = Pt(8)
        r = cap.add_run(figure_caption or "Validated project figure")
        set_run_font(r, size=9, color=MID_GRAY, italic=True)


def add_cover(doc):
    image = PROJECT_ROOT / "frontend" / "web" / "public" / "og.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(24)
    picture = p.add_run().add_picture(str(image), width=Inches(6.5))
    set_picture_alt_text(
        picture,
        "FinAccess Eswatini report cover graphic with the project title and analytical motif.",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("PHASE 1-13 IMPLEMENTATION REPORT")
    set_run_font(r, size=24, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Tasks completed, technical reasoning, evidence, and current project state")
    set_run_font(r, size=12.5, color=MID_GRAY)

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2500, 6860])
    set_table_borders(table, color="D9E3F0")
    entries = (
        ("Project", "FinAccess Eswatini"),
        ("Prepared for", "Project owner and portfolio reviewers"),
        ("Prepared by", "Thando F. Dlamini"),
        ("Reporting date", "11 August 2026"),
    )
    for row, (label, value) in zip(table.rows, entries):
        set_cell_shading(row.cells[0], PALE_BLUE)
        r = row.cells[0].paragraphs[0].add_run(label.upper())
        set_run_font(r, size=8.5, color=BLUE, bold=True)
        r = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(r, size=10.5, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_callout(
        doc,
        "Current position",
        "All thirteen approved phases are complete. Two validated explainable models, a combined prediction API, and the selected Signal web application are publicly available under one Vercel domain.",
        fill=PALE_TEAL,
        accent=TEAL,
    )


def add_executive_summary(doc):
    doc.add_page_break()
    doc.add_heading("Executive summary", level=1)
    add_paragraph(
        doc,
        "FinAccess Eswatini is a portfolio-grade analytical and decision-support proof of concept built from World Bank Global Findex Eswatini microdata. The work was deliberately executed phase by phase so that data understanding, leakage control, statistical reasoning, predictive modelling, explainability, API engineering, and product design could each be validated before the next layer was introduced.",
    )
    add_paragraph(
        doc,
        "The system contains two independent machine-learning engines: one estimates whether a profile is likely to be financially included and the other estimates whether the profile is likely to use mobile money. The models share a product experience, but they do not share feature blueprints, selection decisions, evaluation splits, or explainers.",
    )
    add_callout(
        doc,
        "Outcome",
        "The project has progressed from a raw 1,051-row, 199-variable survey file to two leakage-reviewed model pipelines, validated SHAP explanations, a 17-field combined API, and a responsive Signal frontend with five product areas.",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    doc.add_heading("Project at a glance", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3000, 1860, 4500])
    set_table_borders(table)
    headers = ("Measure", "Result", "Interpretation")
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=9.5, color=WHITE, bold=True)
    rows = (
        ("Raw dataset", "1,051 x 199", "One Eswatini survey extract; respondent-level microdata"),
        ("Weighted financial inclusion", "43.1%", "Financial institution account estimate"),
        ("Weighted mobile money", "50.4%", "Mobile money account estimate"),
        ("Model 1", "ROC-AUC 0.745", "Gradient Boosting on a protected 211-row holdout"),
        ("Model 2", "ROC-AUC 0.726", "Logistic Regression on a protected 210-row holdout"),
        ("Explainability", "2 SHAP explainers", "Five model-derived factors can be returned per outcome"),
        ("API", "17 validated inputs", "One request runs both models and both explanation paths"),
        ("Frontend", "Signal selected", "Overview, two analytics areas, assessment, and methodology"),
        ("Automated validation", "101 Python tests", "Plus 4 deployment API tests, 5 rendered routes, and 15 public checks"),
    )
    for row_values in rows:
        row = table.add_row()
        for idx, text in enumerate(row_values):
            r = row.cells[idx].paragraphs[0].add_run(text)
            set_run_font(r, size=9.2, color=DARK, bold=(idx == 1))
        if len(table.rows) % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F8FAFC")
    repeat_header(table.rows[0])

    doc.add_heading("Governing engineering principles", level=2)
    add_bullets(doc, [
        ("Phase gating. ", "Each phase was completed, validated, reported, and stopped before the next phase began."),
        ("Leakage prevention. ", "Predictors were included only when defensible for the specific target; parallel outcomes and post-outcome behaviours were excluded."),
        ("Separation of claims. ", "Weighted description, statistical association, prediction, and causal inference were treated as different analytical activities."),
        ("Protected evaluation. ", "Identical predictor profiles were grouped so matching profiles could not cross training, validation, or holdout boundaries."),
        ("Reproducibility. ", "Raw data and model artifacts were protected by hashes, reusable modules, executed notebooks, and automated tests."),
        ("Human-readable output. ", "The product leads with plain-language answers, then probabilities and model-derived explanation factors."),
    ])


def add_architecture(doc):
    doc.add_heading("Final system architecture", level=1)
    add_numbered(doc, [
        "World Bank Global Findex Eswatini raw microdata is verified against an immutable file hash.",
        "Model-specific cleaning and deterministic feature engineering create separate final matrices.",
        "Independent preprocessing and estimators produce financial-inclusion and mobile-money probabilities.",
        "Model-matched SHAP explainers generate signed global and individual contributions.",
        "FastAPI validates one 17-field profile, runs both pipelines, and returns structured natural-language results.",
        "The selected Signal frontend presents descriptive evidence, assessment inputs, predictions, and methodology.",
        "Vercel routes the Next.js interface and FastAPI service under one public HTTPS domain.",
    ])
    add_callout(
        doc,
        "Scope boundary",
        "The deployed system is a portfolio proof of concept. It is not a production financial decision, eligibility, creditworthiness, or official World Bank classification system.",
        fill=PALE_TEAL,
        accent=TEAL,
    )


def build_report():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_executive_summary(doc)
    add_architecture(doc)

    add_phase(
        doc, 1, "Project Setup and Data Audit",
        "Establish a trustworthy project foundation and understand the supplied raw data without changing it.",
        [
            "Created the professional repository structure and placed the source CSV under data/raw.",
            "Verified the file encoding, delimiter, SHA-256 fingerprint, dimensions, headers, and row widths.",
            "Confirmed both targets are present, complete, binary, and coded 0/1.",
            "Audited missingness, data types, value sets, exact duplicates, plausible ranges, identifiers, metadata, and constant fields.",
            "Created an executed data-understanding notebook and reproducible audit reports.",
        ],
        [
            ("Immutable raw-data contract", "The source hash prevents accidental replacement or silent mutation from contaminating later phases."),
            ("No early recoding", "Numeric survey codes can represent categories, refusals, routing states, or real values. Their meaning had to be confirmed before transformation."),
            ("Retain duplicate candidates", "Respondent identifiers were unique, so identical profiles were treated as plausible respondents rather than deleted automatically."),
            ("Separate weights from predictors", "The survey weight is useful for population description but is not a personal characteristic available to an inference user."),
        ],
        [
            "Validated shape: 1,051 respondents and 199 variables.",
            "Found 14 fully missing columns, 107 columns at least 50% missing, and no exact duplicate rows in the full file.",
            "Identified wpid_random as unique respondent identifier and six constant metadata/no-variance fields.",
            "Confirmed age, income quintile, targets, weights, gender code set, and employment code set passed basic plausibility checks.",
        ],
        note=("Key risk discovered", "The dataset contains many financial behaviours adjacent to the targets. Using all available variables would create severe leakage and unrealistically high performance.", "FFF0ED", RED),
    )

    add_phase(
        doc, 2, "Data Dictionary and Feature Eligibility",
        "Define every variable and establish defensible, model-specific feature boundaries before cleaning or modelling.",
        [
            "Mapped all 199 variables to official World Bank metadata and created a complete data dictionary.",
            "Classified variables into demographic, socioeconomic, digital, financial, identity, target, identifier, and metadata groups.",
            "Reviewed each variable independently for Model 1 and Model 2 eligibility, leakage risk, missingness, timing, redundancy, and conceptual scope.",
            "Documented direct leakage, post-outcome behaviour, parallel outcomes, identifiers, metadata, all-missing fields, and high-missing exclusions.",
            "Confirmed the 2024 survey year within the Global Findex 2025 publication/database edition.",
        ],
        [
            ("Model-specific eligibility", "Changing only the target would not be defensible because a variable can be pre-outcome for one prediction problem and leakage for the other."),
            ("Conservative leakage policy", "Direct derivatives, the other target, and same-period financial outcomes were excluded unless a strong pre-outcome justification existed."),
            ("Source-backed coding", "The dictionary and DDI were used to interpret special response codes and routing rather than relying on column names or memory of Findex surveys."),
            ("Conditional candidates", "Potentially useful digital variables were carried forward for later routing and engineering review instead of being accepted blindly."),
        ],
        [
            "Documented all 199 variables with complete Model 1 and Model 2 decisions.",
            "Defined 16 Phase 2 candidates for financial inclusion and 26 for mobile money.",
            "Protected account_fin and account_mob as targets and excluded their parallel-outcome use.",
            "Produced leakage review, target review, feature blueprint, and source notes.",
        ],
        note=("Why this phase mattered", "Feature eligibility was treated as a research-design decision, not a convenience step. This is the primary defence against misleading portfolio performance.", PALE_TEAL, TEAL),
    )

    add_phase(
        doc, 3, "Data Cleaning and Preprocessing",
        "Convert approved raw inputs into reproducible, semantically honest modelling datasets without fitting on future holdout data.",
        [
            "Implemented reusable cleaning functions and model-specific preprocessing templates.",
            "Converted routed blanks to Not applicable / skipped and explicit special responses to meaningful categorical states.",
            "Retained age as validated numeric input and treated coded predictors as categories rather than artificial numeric scales.",
            "Generated separate cleaned datasets for account_fin and account_mob with no unresolved missing cells.",
            "Added strict schema, code-set, target, leakage, and unknown-category tests.",
        ],
        [
            ("Semantic missingness", "Questionnaire routing carries information about eligibility and access. Treating routed blanks as random nulls would erase that meaning."),
            ("No silent imputation", "The ten blank education responses became an explicit missing/nonresponse category instead of being replaced with an assumed typical value."),
            ("Unfitted preprocessors", "Preprocessing objects remained templates so category learning and parameter fitting could occur only inside training folds."),
            ("Central feature policy", "The same approved lists drive dictionary generation, cleaning, and later modelling, reducing policy drift."),
        ],
        [
            "Model 1 cleaned matrix: 1,051 rows, 16 predictors, one target.",
            "Model 2 cleaned matrix: 1,051 rows, 26 predictors, one target.",
            "Age range was 15-100, with zero non-integer or out-of-range values.",
            "No identifiers, metadata, parallel targets, unapproved fields, or unresolved null cells entered either output.",
        ],
    )

    add_phase(
        doc, 4, "Exploratory Data Analysis",
        "Build a question-driven, survey-weighted analytical story before formal inference or modelling.",
        [
            "Calculated survey-weighted national estimates and retained unweighted counts for sample transparency.",
            "Compared both outcomes across gender, age, education, income, employment, internet use, phone ownership, and phone type.",
            "Created professional PNG and SVG figures with a complete chart manifest.",
            "Recorded analytical findings using association language and explicit interpretation guardrails.",
            "Embedded saved tables and figures in the executed Phase 4 notebook.",
        ],
        [
            ("Weighted description", "The supplied survey weight makes the national descriptive estimates more appropriate than raw sample proportions."),
            ("Question-led charts", "Visuals were tied to project questions so the EDA communicates an analytical narrative instead of producing a generic chart gallery."),
            ("No premature inference", "Phase 4 describes patterns only; p-values, effect sizes, and models were deliberately reserved for later phases."),
            ("Transparent denominators", "Weighted rates are accompanied by unweighted sample sizes so readers can see the evidence supporting each subgroup estimate."),
        ],
        [
            "Weighted financial-inclusion estimate: 43.1%; weighted mobile-money estimate: 50.4%.",
            "Financial inclusion ranged from 36.8% for primary education or less to 82.4% for tertiary education or more.",
            "Financial inclusion rose from 34.1% in income quintile 1 to 65.0% in quintile 5.",
            "Mobile-money adoption was 60.5% among recent internet users and 39.7% in the combined no/DK/refused group.",
            "Urban/rural analysis was not possible because urbanicity is constant in the extract.",
        ],
        note=("Interpretation boundary", "These weighted comparisons are observational and bivariate. They show association, not causation, and do not adjust for confounding.", "FFF7E6", GOLD),
        figure=PROJECT_ROOT / "reports" / "phase_4" / "figures" / "01_overall_access_rates.png",
        figure_caption="Figure 1. Survey-weighted and unweighted access estimates from the validated Phase 4 analysis.",
    )

    add_phase(
        doc, 5, "Statistical Analysis",
        "Test the principal EDA associations formally while controlling false discoveries and preserving non-causal interpretation.",
        [
            "Pre-specified eight tests per outcome: seven categorical tests and one numeric age comparison.",
            "Applied chi-square tests, bias-corrected Cramer's V, Mann-Whitney tests, and rank-biserial effect sizes as appropriate.",
            "Controlled multiple testing with Benjamini-Hochberg FDR separately within each outcome.",
            "Checked expected-count assumptions for all categorical tests.",
            "Produced complete result, contingency, age-distribution, and summary tables.",
        ],
        [
            ("Pre-specification", "Defining tests before inspecting significance reduces selective reporting and makes the analytical scope auditable."),
            ("Effect sizes with p-values", "Statistical significance alone does not show practical magnitude, so association strength was reported alongside adjusted significance."),
            ("Unweighted inference", "Without survey strata and cluster variables, the project avoided implying design-corrected survey inference. Weighted rates remain descriptive context."),
            ("FDR control", "Multiple comparisons increase false-positive risk; Benjamini-Hochberg provides a transparent correction without treating every test as isolated."),
        ],
        [
            "Seven of eight tests remained significant after FDR adjustment for each target; gender did not.",
            "All 14 categorical tests passed the expected-count rule; the minimum expected cell count was 36.35.",
            "Largest financial-inclusion categorical effect: income quintile, bias-corrected Cramer's V = 0.270.",
            "Largest mobile-money categorical effect: phone type, bias-corrected Cramer's V = 0.236.",
            "Age effects were small: rank-biserial 0.203 for inclusion and 0.108 for mobile money.",
        ],
    )

    add_phase(
        doc, 6, "Feature Engineering",
        "Create interpretable, deterministic, leakage-safe predictors and freeze final model-specific matrices.",
        [
            "Created fixed age groups and a phone-access tier for both models.",
            "Created internet-engagement and data-purchase patterns for Model 2 only.",
            "Reviewed every engineered proposal for target use, observation timing, interpretability, and redundancy.",
            "Rejected online-activity breadth and an arbitrary digital-access score.",
            "Generated final modelling matrices and documented transformation specifications.",
        ],
        [
            ("Interpretability first", "Derived features map to concepts a user can understand and an explanation layer can name clearly."),
            ("Target-independent derivation", "No target values or learned model parameters were used to create engineered predictors."),
            ("Model 2 specificity", "Internet frequency and data-purchase routing are more directly relevant to mobile-money adoption and were not forced into Model 1."),
            ("Reject arbitrary composites", "A digital score would hide heterogeneous behaviours, impose subjective weights, and make leakage review and explanation less transparent."),
        ],
        [
            "Final Model 1 matrix: 1,051 rows and 15 categorical predictors.",
            "Final Model 2 matrix: 1,051 rows and 16 categorical predictors.",
            "Both final matrices contain zero null cells and preserve target distributions.",
            "No split or model training occurred, preserving the dedicated modelling phase boundary.",
        ],
        note=("Carried-forward caution", "Recent digital behaviours for Model 2 overlap the target observation period. They were retained with explicit documentation and remain a conceptual limitation.", "FFF7E6", GOLD),
    )

    add_phase(
        doc, 7, "Model 1: Financial Inclusion",
        "Train and select the financial-inclusion classifier using protected, group-aware evaluation.",
        [
            "Compared Logistic Regression, Decision Tree, Random Forest, and Gradient Boosting pipelines.",
            "Created a deterministic 80/20 group-aware split and five-fold StratifiedGroupKFold training evaluation.",
            "Tuned candidates using training data only and selected with a one-standard-error complexity rule.",
            "Evaluated the selected pipeline once on the protected holdout, including discrimination, classification, calibration, and bootstrap intervals.",
            "Saved the complete preprocessing-plus-estimator pipeline and integrity metadata.",
        ],
        [
            ("Group identical profiles", "Matching predictor profiles must not appear in both training and evaluation sets because that would reward memorisation and inflate generalisation estimates."),
            ("One-standard-error rule", "A candidate had to be statistically competitive with the best mean cross-validation AUC; complexity was then considered rather than selecting a noisy maximum automatically."),
            ("Holdout used once", "The protected test set was not used for tuning or model choice, preserving an honest final estimate."),
            ("Pipeline persistence", "Saving preprocessing with the estimator ensures training and inference apply exactly the same transformations."),
        ],
        [
            "Selected Gradient Boosting: 200 estimators, learning rate 0.03, maximum depth 2, minimum leaf size 5.",
            "Holdout metrics: ROC-AUC 0.745, accuracy 0.706, precision 0.717, recall 0.704, F1 0.710.",
            "Confusion matrix: TN 73, FP 30, FN 32, TP 76.",
            "Profile overlap was zero in the holdout and every cross-validation fold.",
            "Reloaded pipeline reproduced predictions; SHA-256 integrity check passed.",
        ],
        note=("Selection note", "The 0.50 classification threshold remains provisional. Model quality is suitable for a portfolio proof of concept, not a production decision authority.", "FFF7E6", GOLD),
    )

    add_phase(
        doc, 8, "Model 2: Mobile Money Adoption",
        "Run a fully independent modelling workflow for mobile-money adoption and protect the completed Model 1 artifact.",
        [
            "Repeated the candidate-model workflow independently with the Model 2 feature blueprint and separate random seeds.",
            "Used a 210-row protected holdout and group-aware five-fold cross-validation on the training partition.",
            "Compared Logistic Regression, Decision Tree, Random Forest, and Gradient Boosting candidates.",
            "Selected, evaluated, saved, and reloaded the complete Model 2 pipeline.",
            "Verified the Model 1 pipeline hash remained unchanged throughout Phase 8.",
        ],
        [
            ("Independent workflow", "Mobile-money adoption has a different conceptual boundary, final feature matrix, duplicate-profile structure, split, candidate results, and explanation needs."),
            ("Prefer defensible simplicity", "Random Forest produced the highest mean CV AUC, but Logistic Regression was within the one-standard-error eligibility range and had the lowest predefined complexity tier."),
            ("Model 1 safeguard", "Hash verification prevented work on Model 2 from silently replacing or altering the already validated financial-inclusion pipeline."),
            ("Balanced class treatment", "The selected logistic model used balanced class weights to reduce majority-class dominance during fitting."),
        ],
        [
            "Selected Logistic Regression with C = 0.05 and balanced class weights.",
            "Holdout metrics: ROC-AUC 0.726, accuracy 0.676, precision 0.721, recall 0.721, F1 0.721.",
            "Confusion matrix: TN 54, FP 34, FN 34, TP 88.",
            "Profile overlap was zero in the holdout and every cross-validation fold.",
            "The saved Model 1 hash before and after Phase 8 was identical.",
        ],
    )

    add_phase(
        doc, 9, "Explainability",
        "Create faithful global and individual explanations for both validated pipelines.",
        [
            "Built a TreeExplainer for the Gradient Boosting financial-inclusion pipeline.",
            "Built a LinearExplainer for the Logistic Regression mobile-money pipeline.",
            "Calculated global mean absolute SHAP importance on each protected holdout.",
            "Aggregated one-hot encoded contributions back to recognisable source variables.",
            "Generated individual examples at low, boundary, and high predicted probabilities with five ranked factors each.",
            "Saved model-matched explainers and reusable inference explanation logic.",
        ],
        [
            ("Model-matched explainers", "Tree and linear estimators have different mathematical structures, so each received an appropriate SHAP implementation."),
            ("Source-level aggregation", "Users should see education level or age group, not fragmented one-hot column names. Aggregation preserves signed additivity while improving comprehension."),
            ("Faithfulness before wording", "Human-readable reasons are generated from signed SHAP contributions; drivers are not invented manually."),
            ("Holdout-based global importance", "Global explanations are evaluated on protected observations rather than the training data used to fit the models."),
        ],
        [
            "Explained all 211 Model 1 holdout rows and all 210 Model 2 holdout rows.",
            "Model 1 leading factors: age group, workforce status, income quintile, education, and recent internet use.",
            "Model 2 leading factors: SIM registration, age group, internet engagement, data-purchase pattern, and income quintile.",
            "Maximum probability reconstruction error was 3.33 x 10^-16.",
            "Both explainers reloaded successfully and both pipeline hashes remained unchanged.",
        ],
        note=("Interpretation boundary", "SHAP describes how a fitted model distributes prediction differences relative to its baseline. It does not establish causation or policy impact.", "FFF7E6", GOLD),
        figure=PROJECT_ROOT / "reports" / "phase_9" / "figures" / "01_global_shap_importance.png",
        figure_caption="Figure 2. Global source-feature SHAP importance for both validated models on their protected holdouts.",
    )

    add_phase(
        doc, 10, "Prediction API",
        "Expose both validated models and explanation paths through one strict, production-style assessment service.",
        [
            "Built a FastAPI service with health, combined assessment, OpenAPI, and interactive documentation endpoints.",
            "Defined a strict 17-field Pydantic request schema with permitted categories and cross-field routing checks.",
            "Loaded both complete pipelines and their matching explainers with artifact-hash verification.",
            "Returned two natural-language answers, probabilities, provisional thresholds, five explanation factors per outcome, warnings, and a responsible-use disclaimer.",
            "Added structured error handling and sanitized model/artifact failure responses.",
        ],
        [
            ("One profile, two results", "The product should feel like one financial-access platform; a user enters relevant characteristics once while each model applies its own preprocessing."),
            ("Strict validation", "Rejecting extra, missing, invalid, or contradictory inputs prevents undefined inference behaviour and makes the API contract explicit."),
            ("Hash verification", "A pipeline and explainer must remain a validated pair. Integrity checks prevent mismatched explanation artifacts from loading silently."),
            ("Human question first", "The API answers whether the profile is likely to be included or use mobile money before presenting supporting percentages."),
        ],
        [
            "Eight validation scenarios passed, including invalid categories, missing/extra fields, contradictory routing, and unavailable artifacts.",
            "A validated example returned 26.9% financial-inclusion likelihood and 36.8% mobile-money likelihood, with five factors for each.",
            "Health checks confirmed both model/explainer pairs were ready.",
            "Submitted profiles are not persisted; CORS and public hosting remained outside Phase 10.",
        ],
    )

    add_phase(
        doc, 11, "Web Application",
        "Create a polished, responsive product experience for analysis, assessment, explainability, and technical methodology.",
        [
            "Developed three complete light-interface directions: The Ledger, Open Field, and Signal.",
            "Implemented five product areas in the shared experience: Overview, Financial Inclusion, Mobile Money, Financial Access Assessment, and Methodology.",
            "Integrated the 17-field assessment with a same-origin server proxy to the Phase 10 API.",
            "Selected Signal as the final direction while retaining the other concepts as archived design history.",
            "Refined the Signal homepage to lead with Financial Access in Eswatini, Start an assessment, Review the evidence, and a developer credit for Thando F. Dlamini.",
            "Created a server-free offline visual review package for reliable local review.",
        ],
        [
            ("Multiple concepts before selection", "Visual direction is a product decision. Building three complete variants allowed evidence-based selection without changing analytical functionality."),
            ("Shared product core", "All variants use the same content, data, assessment, and explanation components, avoiding divergent functionality and duplicated maintenance."),
            ("Signal selection", "Signal offers a clean light fintech presentation with pale blue, cobalt, and teal while preserving an analytical, portfolio-ready tone."),
            ("Same-origin proxy", "The frontend calls its own server route, which then reaches FastAPI. This keeps the browser contract simple and defers production CORS decisions to deployment."),
            ("Offline review fallback", "A self-contained HTML review was added after the local production preview failed to serve a stylesheet asset reliably. The fallback removes server dependence for visual approval."),
        ],
        [
            "Production build and frontend lint passed.",
            "Five rendered-route tests passed for the selected homepage, archived variants, and social metadata.",
            "All 15 frontend contract checks passed, including accessibility, responsiveness, selected content, API integration, and phase-scope controls.",
            "The full repository regression suite passed 89 Python tests.",
            "Deployment has not started; Phase 12 remains the explicit next phase.",
        ],
    )

    add_phase(
        doc, 12, "Deployment",
        "Publish the selected frontend and both-model inference API as one public, recruiter-accessible application.",
        [
            "Packaged the Signal Next.js frontend and FastAPI inference service as separate Vercel Services under one project.",
            "Routed all browser assessment traffic to same-origin /api endpoints and removed the cross-host Render dependency.",
            "Published all six validated pipeline, metadata, and model-matched explainer artifacts while excluding respondent microdata.",
            "Verified public unauthenticated access, health, OpenAPI, interactive documentation, valid prediction, and invalid-input behaviour.",
            "Compared the production response with the validated Phase 10 probabilities and every returned SHAP factor.",
        ],
        [
            ("One public domain", "A recruiter should reach the product without a platform login, cross-host proxy, or browser CORS dependency."),
            ("Artifact integrity", "The deployed service verifies pipeline and explainer hashes before it reports healthy, preventing silent model/explainer mismatch."),
            ("Publication minimisation", "Only inference artifacts and application code are deployed; raw and processed respondent records remain local."),
            ("Live equivalence", "A deployment is not considered valid merely because it responds. Its probabilities and explanation factors must match the approved local reference."),
        ],
        [
            "All 15 public deployment checks passed.",
            "The application, API information, health endpoint, interactive documentation, and combined assessment are available at finaccess-eswatini.vercel.app.",
            "The deployment package passed its API suite; the frontend passed lint, production build, dependency audit, and five rendered-route checks.",
            "No public authentication, respondent microdata, live environment file, private key, or external Render API dependency is required.",
        ],
        note=("Deployment note", "Vercel Services and the Python runtime remain platform dependencies to regression-test. Automatic Git deployments require Vercel GitHub App access; the validated release used the authenticated CLI.", "FFF7E6", GOLD),
    )

    add_phase(
        doc, 13, "Final Portfolio Polish",
        "Make the complete analytical, engineering, and product work immediately understandable and reproducible for portfolio reviewers.",
        [
            "Rewrote the analytical and deployment READMEs around the live product, verified results, architecture, reproduction path, and responsible-use limitations.",
            "Captured four fresh screenshots directly from the public application, including a real two-model assessment response.",
            "Added a reproducible Chrome DevTools screenshot script and mirrored selected images into the deployment repository.",
            "Added and executed the thirteenth phase-aligned notebook with saved model results and live product visuals.",
            "Extended this implementation-and-rationale document through deployment and final portfolio polish.",
            "Audited both repositories for publishable secrets, respondent microdata, caches, runtime output, and generated build artifacts.",
        ],
        [
            ("Reviewer-first ordering", "The final README leads with the product and evidence because a portfolio reviewer should understand the problem, output, and technical depth within seconds."),
            ("Authentic screenshots", "Images were generated from the live public release rather than supplied mockups, making them reproducible evidence of the deployed product."),
            ("Separate repositories", "The analytical workflow and deployable product remain independently versioned so publication scope and respondent-data safeguards stay explicit."),
            ("Machine-checked polish", "Screenshots, notebooks, documentation, repository safety, and the live product are tested as deliverables rather than treated as cosmetic extras."),
        ],
        [
            "Four live PNG screenshots passed dimension and non-blank-image checks and were visually inspected.",
            "All 13 notebooks contain executed code cells, saved outputs, and no saved execution errors.",
            "Both publishable repositories contain no secret file, raw respondent data file, cache, runtime directory, dependency directory, or generated build directory.",
            "The production regression gate passed all 15 live checks after portfolio changes.",
        ],
        note=("Owner-controlled publication", "The live application is public and requires no sign-in. Source-repository visibility remains private until the project owner explicitly approves making the GitHub repositories public.", "FFF7E6", GOLD),
        figure=PROJECT_ROOT / "reports" / "phase_13" / "screenshots" / "01_overview.png",
        figure_caption="Figure 3. Final Signal overview captured from the live public Vercel application during Phase 13.",
    )

    cross_phase_heading = doc.add_heading("Cross-phase decisions and rationale", level=1)
    cross_phase_heading.paragraph_format.page_break_before = True
    decisions = (
        ("Raw microdata remains local and Git-ignored", "The file contains respondent-level data and publication/licensing conditions require a deliberate release decision."),
        ("Survey weights support description, not personal prediction inputs", "Weights adjust sample representation; they are not stable personal characteristics available in an assessment."),
        ("Both targets receive independent feature and model workflows", "Financial institution accounts and mobile-money accounts are related but distinct outcomes with different leakage boundaries."),
        ("Categorical semantics are preserved", "Education, income quintile, routing, and nonresponse are not treated as continuous distances or arbitrary numeric quantities."),
        ("Identical profiles are grouped during evaluation", "Preventing profile overlap reduces optimistic performance caused by repeated feature patterns."),
        ("Accuracy is not the selection criterion", "ROC-AUC, F1, calibration, uncertainty, overfitting, complexity, and interpretability all inform selection."),
        ("Explanations come from SHAP", "Displayed factors correspond to actual signed model contributions rather than generic narratives."),
        ("Thresholds remain provisional", "A 0.50 cutoff is operationally simple but has not been optimised against a real decision cost or public policy objective."),
    )
    for label, detail in decisions:
        add_paragraph(doc, f"{label}. {detail}", bold_prefix=f"{label}. ")

    doc.add_heading("Current risks and limitations", level=1)
    add_bullets(doc, [
        "The project is observational and cannot establish that any characteristic causes financial inclusion or mobile-money adoption.",
        "The country extract has no usable urban/rural variation because urbanicity is constant.",
        "Survey-design-corrected inference was not possible without strata and cluster variables.",
        "Small protected holdouts create material uncertainty around performance estimates.",
        "Some identical predictor profiles have conflicting targets, placing an irreducible ceiling on deterministic profile-level prediction.",
        "Recent digital-behaviour variables in Model 2 overlap the target observation period and remain a conceptual limitation.",
        "The assessment threshold of 0.50 is provisional and not tied to a real-world cost-benefit policy.",
        "The application is a portfolio proof of concept, not a nationwide decision engine, eligibility tool, or official World Bank classification.",
        "Vercel Services and the Python runtime are platform dependencies that require regression monitoring.",
        "Both GitHub repositories remain private until the project owner explicitly approves public source visibility.",
    ])

    doc.add_heading("Roadmap completion", level=1)
    add_paragraph(doc, "The approved thirteen-phase roadmap is complete. Any future work should be treated as a separately approved extension with its own scope, validation, and responsible-use review.")
    add_status_table(doc, "COMPLETE WITH NOTES", "Public analytical product, two independently validated model pipelines, SHAP explanations, FastAPI service, Next.js interface, deployment evidence, and recruiter-facing documentation.")

    doc.add_heading("Primary evidence and deliverables", level=1)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [850, 3350, 5160])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Phase", "Notebook", "Primary report or artifact")):
        set_cell_shading(cell, NAVY)
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=9.5, color=WHITE, bold=True)
    evidence_rows = (
        ("1", "01_data_understanding.ipynb", "reports/phase_1/data_quality_report.md"),
        ("2", "02_data_dictionary_feature_eligibility.ipynb", "reports/phase_2/feature_blueprint.md"),
        ("3", "03_data_cleaning_preprocessing.ipynb", "reports/phase_3/phase3_summary.md"),
        ("4", "04_exploratory_analysis.ipynb", "reports/phase_4/eda_report.md"),
        ("5", "05_statistical_analysis.ipynb", "reports/phase_5/statistical_analysis_report.md"),
        ("6", "06_feature_engineering.ipynb", "reports/phase_6/feature_engineering_report.md"),
        ("7", "07_financial_inclusion_model.ipynb", "reports/phase_7/model1_report.md"),
        ("8", "08_mobile_money_model.ipynb", "reports/phase_8/model2_report.md"),
        ("9", "09_model_explainability.ipynb", "reports/phase_9/explainability_report.md"),
        ("10", "10_prediction_api.ipynb", "reports/phase_10/api_report.md"),
        ("11", "11_web_application.ipynb", "reports/phase_11/web_application_report.md"),
        ("12", "12_deployment.ipynb", "reports/phase_12/deployment_report.md"),
        ("13", "13_portfolio_polish.ipynb", "reports/phase_13/final_portfolio_report.md"),
    )
    for values in evidence_rows:
        row = table.add_row()
        for idx, text in enumerate(values):
            r = row.cells[idx].paragraphs[0].add_run(text)
            set_run_font(r, size=8.6, color=DARK, bold=(idx == 0))
        if int(values[0]) % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F8FAFC")
    repeat_header(table.rows[0])

    doc.add_heading("Source note", level=1)
    add_paragraph(doc, "Dataset reference: World Bank Global Findex, Eswatini reference SWZ_2024_FINDEX_v02_M. DOI: https://doi.org/10.48529/5rsc-p773. Microdata catalog: https://microdata.worldbank.org/catalog/7900.")
    add_paragraph(doc, "The 2024 survey year and the Global Findex 2025 publication/database edition are both retained in project documentation to avoid misrepresenting the source.")
    add_callout(doc, "Document status", "This report documents the repository state after completion and validation of all thirteen approved phases on 11 August 2026.", PALE_TEAL, TEAL)

    doc.core_properties.title = "FinAccess Eswatini - Phase 1-13 Implementation Report"
    doc.core_properties.subject = "Completed tasks, technical rationale, validation evidence, and project status"
    doc.core_properties.author = "Thando F. Dlamini"
    doc.core_properties.keywords = "FinAccess Eswatini, financial inclusion, mobile money, machine learning, explainability"
    doc.core_properties.comments = "Professional project documentation generated from validated Phase 1-13 repository artifacts."
    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    build_report()
